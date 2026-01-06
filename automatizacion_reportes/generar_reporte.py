import polars as pl
from docx import Document
from datetime import datetime


def generar_reporte():
    df = pl.read_csv("data/ventas_limpio.csv", separator=";")

    total_registros = df.height
    total_ventas = df["monto"].sum()

    resumen = (
        df.group_by("región")
        .agg([
            pl.col("monto").sum().alias("total_ventas"),
            pl.col("cliente").n_unique().alias("clientes_unicos")
        ])
    )

    # Guardar resumen CSV (opcional, interno)
    resumen.write_csv(
        "data/resumen.csv",
        separator=";",
        encoding="utf8"
    )

    # Documento Word
    doc = Document()
    doc.add_heading("Reporte automático de ventas", level=1)

    fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
    doc.add_paragraph(f"Generado automáticamente el {fecha}")

    doc.add_heading("Resumen general", level=2)
    doc.add_paragraph(f"Total de registros procesados: {total_registros}")
    doc.add_paragraph(f"Total de ventas: ${total_ventas:,.2f}")

    doc.add_heading("Ventas por región", level=2)

    tabla = doc.add_table(rows=1, cols=3)
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "Región"
    encabezados[1].text = "Total Ventas"
    encabezados[2].text = "Clientes Únicos"

    for fila in resumen.iter_rows():
        celdas = tabla.add_row().cells
        celdas[0].text = fila[0]
        celdas[1].text = f"${fila[1]:,.2f}"
        celdas[2].text = str(fila[2])

    doc.save("reporte_ventas.docx")
