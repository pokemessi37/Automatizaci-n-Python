from docx import Document
import polars as pl


def generar_reporte_word():
    # Leer CSV limpio (forzamos separador y encoding)
    df = pl.read_csv(
        "data/ventas_limpio.csv",
        separator=";",
        encoding="utf-8"
    )

    # Normalizar nombres de columnas (por las dudas)
    df = df.rename({
        col: (
            col.strip()
            .lower()
            .replace("á", "a")
            .replace("é", "e")
            .replace("í", "i")
            .replace("ó", "o")
            .replace("ú", "u")
        )
        for col in df.columns
    })

    # Validaciones mínimas (evita errores silenciosos)
    columnas_requeridas = {"cliente", "region", "monto"}
    faltantes = columnas_requeridas - set(df.columns)

    if faltantes:
        raise ValueError(
            f"El CSV no contiene las columnas necesarias: {', '.join(faltantes)}"
        )

    # Crear documento Word
    doc = Document()
    doc.add_heading("Reporte de Ventas", level=1)

    doc.add_paragraph(
        "Este reporte fue generado automáticamente a partir de los datos procesados."
    )

    doc.add_heading("Resumen por región", level=2)

    # Crear tabla
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Light List"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Región"
    hdr[1].text = "Total Ventas"
    hdr[2].text = "Clientes Únicos"

    # Agrupar y calcular métricas
    resumen = (
        df.group_by("region")
        .agg(
            pl.col("monto").sum().alias("total_ventas"),
            pl.col("cliente").n_unique().alias("clientes_unicos")
        )
        .sort("total_ventas", descending=True)
    )

    # Cargar datos en la tabla
    for region, total, clientes in resumen.iter_rows():
        row = tabla.add_row().cells
        row[0].text = str(region)
        row[1].text = f"$ {total:,.2f}"
        row[2].text = str(clientes)

    # Guardar archivo
    doc.save("data/reporte_ventas.docx")
